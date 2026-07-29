from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "EMI_Flow_FDE_Submission.docx"
ARCH = ROOT / "submission" / "emi_architecture.png"

BLUE = "2E74B5"
DARK_BLUE = "16324F"
INK = "172B24"
GREEN = "1E7654"
LIGHT_GREEN = "E8F4EE"
GOLD = "A96E16"
LIGHT_GOLD = "FFF4DE"
RED = "9B1C1C"
LIGHT_RED = "FCEAEA"
MUTED = "61716B"
LIGHT_BLUE = "E8EEF5"
BLACK = "000000"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, italic=False, color=BLACK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("EMI FLOW  |  FDE SUBMISSION EVIDENCE")
    set_run(r, size=8.5, bold=True, color=MUTED)
    p.paragraph_format.space_after = Pt(0)

    fp = section.footer.paragraphs[0]
    left = fp.add_run("Vaishnavi Venkateswaran  |  28 July 2026")
    set_run(left, size=8.5, color=MUTED)
    fp.add_run("\t")
    add_page_number(fp)
    return doc


def add_callout(doc, label, text, status="info"):
    fill = {"pass": LIGHT_GREEN, "partial": LIGHT_GOLD, "gap": LIGHT_RED, "info": LIGHT_BLUE}[status]
    accent = {"pass": GREEN, "partial": GOLD, "gap": RED, "info": BLUE}[status]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(label), size=9, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(text), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(95)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("FDE SUBMISSION REPORT"), size=10, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("EMI Flow"), size=31, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    set_run(p.add_run("EMI Payment Collection Agent"), size=17, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    set_run(
        p.add_run(
            "Architecture, implementation, mandatory tests, demonstration evidence, limitations, and production-readiness plan"
        ),
        size=11,
        italic=True,
        color=MUTED,
    )

    add_callout(
        doc,
        "MY DEMONSTRATION",
        "I built and deployed the application lifecycle end to end using an explicitly labeled mock Gnani trigger plus authenticated simulated post-call webhooks. I also completed and recorded a real multi-turn Gnani Agent Console conversation before my Agent Console call credits were exhausted. Provider-triggered outbound calling remains dependent on Gnani-issued trigger credentials and credits.",
        "partial",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    set_run(p.add_run("Prepared by Vaishnavi Venkateswaran"), size=11, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("28 July 2026"), size=10, color=MUTED)
    doc.add_page_break()


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    header = table.rows[0]
    header.cells[0].text = "Item"
    header.cells[1].text = "Value"
    format_header(header)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        format_body_row(cells, bold_first=True)
    return table


def format_header(row):
    set_repeat_table_header(row)
    prevent_row_split(row)
    for cell in row.cells:
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run(r, size=9, bold=True, color=DARK_BLUE)


def status_color(status):
    return GREEN if status == "PASS" else GOLD if status == "PARTIAL" else RED


def format_body_row(cells, bold_first=False, status_index=None):
    for index, cell in enumerate(cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.08
            for r in p.runs:
                color = status_color(r.text) if status_index == index else INK
                set_run(
                    r,
                    size=8.7,
                    bold=(bold_first and index == 0) or status_index == index,
                    color=color,
                )


def add_matrix(doc, headers, widths, rows, status_index=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    for index, value in enumerate(headers):
        header.cells[index].text = value
    format_header(header)
    for row in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(row):
            cells[index].text = str(value)
        format_body_row(cells, status_index=status_index)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_figure(doc, path, caption, note=None, width=6.35):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    set_run(p.add_run(caption), size=12, bold=True, color=DARK_BLUE)
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        set_run(p.add_run(note), size=9.5, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(width))


def build_architecture():
    canvas = Image.new("RGB", (1800, 920), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    f_title = ImageFont.truetype(bold_path, 34)
    f_body = ImageFont.truetype(font_path, 27)
    f_small = ImageFont.truetype(font_path, 23)

    def box(x, y, w, h, title, lines, fill, outline):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=24, fill=fill, outline=outline, width=4)
        draw.text((x + 26, y + 22), title, font=f_title, fill=outline)
        yy = y + 80
        for line in lines:
            draw.text((x + 26, yy), line, font=f_small, fill=(30, 45, 40))
            yy += 36

    def arrow(x1, y1, x2, y2, label):
        draw.line((x1, y1, x2, y2), fill=(61, 96, 85), width=6)
        draw.polygon([(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)], fill=(61, 96, 85))
        draw.text(((x1 + x2) // 2 - 90, y1 - 42), label, font=f_small, fill=(61, 96, 85))

    box(55, 110, 360, 210, "Client / Dashboard", ["Complete customer data", "Mock trigger form", "Call and outcome views"], (235, 246, 240), (30, 118, 84))
    box(575, 110, 360, 210, "FastAPI", ["Initial_Message", "Post-call webhook", "Validation + idempotency"], (232, 238, 245), (46, 116, 181))
    box(1095, 110, 360, 210, "MongoDB Atlas", ["emi_flow.calls", "Pending + completed", "Transcript + outcome"], (235, 246, 240), (30, 118, 84))
    box(575, 560, 360, 210, "Gnani Console", ["Live voice test evidence", "Evon LLM + Timbre TTS", "Trigger API unavailable"], (255, 244, 222), (169, 110, 22))
    arrow(415, 215, 575, 215, "HTTPS REST")
    arrow(935, 215, 1095, 215, "TLS")
    draw.line((755, 560, 755, 320), fill=(61, 96, 85), width=6)
    draw.polygon([(755, 320), (745, 338), (765, 338)], fill=(61, 96, 85))
    draw.text((785, 405), "initial greeting / post-call", font=f_small, fill=(61, 96, 85))
    draw.text(
        (55, 840),
        "Demo mode replaces the unavailable provider trigger with a mock provider ID; later API, persistence, and dashboard behavior remains real.",
        font=f_body,
        fill=(23, 50, 79),
    )
    canvas.save(ARCH)


def main():
    build_architecture()
    doc = configure_doc()
    add_title_page(doc)

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "I built and deployed EMI Flow as a FastAPI, React, and MongoDB Atlas solution for capturing EMI collection call requests, storing pending call state, processing authenticated and idempotent post-call outcomes, and displaying stage codes, disposition reasons, transcripts, and operational totals."
    )
    add_kv_table(
        doc,
        [
            ("Live dashboard", "https://emi-flow-dashboard.onrender.com"),
            ("Live API", "https://emi-flow-api.onrender.com"),
            ("API documentation", "https://emi-flow-api.onrender.com/docs"),
            ("Source repository", "https://github.com/vaishnaviikv/EMI-payment-collection-agent"),
            ("Gnani agent ID", "bce1ae6dde324dca8aba045f2c86f1fc"),
            ("Live conversation ID", "9b8a1885-efd3-4d9a-9fdc-367de019e116"),
            ("Database", "MongoDB Atlas - emi_flow.calls"),
            ("Deployment", "Render Web Service + Render Static Site"),
            ("Free-tier demo note", "My Render services can spin down after inactivity. Before a live demonstration, I deploy the latest commit to restore a fresh live state, then refresh the dashboard; the first backend request may still take about 50 seconds."),
        ],
    )
    doc.add_heading("Evidence legend", level=2)
    add_callout(doc, "PASS", "Implemented and supported by captured evidence, automated tests, or visible deployed behavior.", "pass")
    add_callout(doc, "PARTIAL", "A meaningful portion is demonstrated, but exact provider behavior or the required named model/version is not fully evidenced.", "partial")

    doc.add_heading("2. Architecture and lifecycle", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ARCH), width=Inches(6.35))
    p = doc.add_paragraph("Figure 1. Deployed architecture and documented mock-provider boundary.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_run(r, size=9, italic=True, color=MUTED)
    add_bullet(doc, "Initial request: the dashboard or client sends complete customer and EMI data to POST /api/Initial_Message.")
    add_bullet(doc, "Persistence: FastAPI validates and inserts a pending call record in MongoDB Atlas.")
    add_bullet(doc, "Provider boundary: mock mode returns a deterministic mock provider ID because the current account exposes no outbound-trigger URL or credentials.")
    add_bullet(doc, "Post-call: an authenticated webhook applies transcript, outcome analytics, stage code, and disposition reason atomically.")
    add_bullet(doc, "Presentation: the dashboard reads persisted records through FastAPI and supports search and outcome filters.")

    doc.add_heading("3. Demonstration expectations", level=1)
    expectations = [
        ("1", "Explain overall architecture", "PASS", "Architecture diagram, deployed URLs, API/Mongo/dashboard lifecycle."),
        ("2", "Show Gnani Voice Agent configuration", "PASS", "Agent details, languages, Evon model, Timbre voice, analytics fields, conversation logs."),
        ("3", "Confirm Prisma ASR, Timbre 2.5 TTS, Evon LLM", "PARTIAL", "Prisma v2.5 is visible in the Gnani API playground and Evon v2.0 Fast is visible in agent settings. The agent uses Timbre G v1.0 because it was the only TTS model available in this Agent Console account; Timbre 2.5 was not offered. The agent Transcriber selection was not captured."),
        ("4", "Initiate a call using FastAPI", "PASS", "Create test call posts complete FDE data; FastAPI validates, stores, and returns triggered status plus mock provider ID."),
        ("5", "Demonstrate multi-turn conversation", "PASS", "Real Gnani console recording and transcript show a 1m08s multi-turn voice test; simulated records add scenario-specific conversations."),
        ("6", "Show post-call webhook received", "PARTIAL", "Authenticated simulated webhook is proven through completed records and webhook_ids. Exact automatic Gnani post-call JSON was not captured."),
        ("7", "Show database record updated", "PASS", "MongoDB screenshot shows completed state, webhook ID, transcript, outcome, disposition, and provider ID."),
        ("8", "Show stage code and disposition reason", "PASS", "Dashboard call detail and mandatory scenario overview show normalized outcomes and reason."),
        ("9", "Demonstrate failure handling", "PASS", "Duplicate delivery, invalid request, missing-call recovery, mocked trigger failure, and timeout paths are implemented and automated."),
        ("10", "Explain production readiness", "PASS", "Section 11 covers credentials, provider contract, security, observability, retries, privacy, and deployment hardening."),
    ]
    add_matrix(doc, ["#", "Expectation", "Status", "Evidence / qualification"], [500, 2550, 900, 5410], expectations, status_index=2)

    doc.add_heading("4. Mandatory test scenarios", level=1)
    scenarios = [
        ("1", "Customer commits to paying today", "PTP_TODAY", "PASS", "Asha Today - Promise To Pay Today"),
        ("2", "Customer provides a future PTP date", "PTP_FUTURE", "PASS", "Ben Future - Promise To Pay"),
        ("3", "Customer states payment is complete", "ALREADY_PAID", "PASS", "Carla Paid - Already Paid"),
        ("4", "Customer requests callback", "CALLBACK_SCHEDULED", "PASS", "Deepak Callback - Callback Scheduled"),
        ("5", "Refuses due to financial difficulty", "RTP_FINANCIAL", "PASS", "Elena Hardship - Financial Hardship"),
        ("6", "Disputes EMI amount", "DISPUTE_CHARGES", "PASS", "Farah Dispute - Charge Dispute"),
        ("7", "Third party answers", "THIRD_PARTY", "PASS", "Gita Account - Third Party"),
        ("8", "Changes language", "PTP_FUTURE", "PASS", "Hari Bilingual - Hindi request plus future PTP"),
        ("9", "Disconnected, no clear outcome", "DSCN", "PASS", "Iris Disconnected - Disconnected"),
        ("10", "Duplicate post-call webhook", "Idempotent duplicate", "PASS", "Same X-Webhook-Id returns duplicate; no second update"),
        ("11", "Invalid initial request", "HTTP 422", "PASS", "Serialization defect found during live run and fixed in commit df76400"),
        ("12", "Gnani trigger fails or times out", "HTTP 502 / 504", "PASS", "Mock failure/timeout store trigger_failed; parametrized API test"),
    ]
    add_matrix(doc, ["#", "Scenario", "Expected", "Status", "Proof"], [420, 2930, 1570, 800, 3640], scenarios, status_index=3)
    add_callout(
        doc,
        "TEST EXECUTION NOTE",
        "The first mandatory-suite run completed scenarios 1-10 and exposed an error-serialization defect during scenario 11. The defect was corrected in commit df76400; the runner was also hardened to retain non-JSON provider responses. The dashboard overview visibly contains all scenario 1-9 records.",
        "info",
    )

    doc.add_heading("5. Acceptance criteria", level=1)
    acceptance = [
        ("FastAPI initiation", "PASS", "Complete FDE request accepted, validated, stored, and returned as triggered in mock mode."),
        ("Customer receives call from configured Gnani agent", "PASS", "Verified live call, recording, and transcript from the configured Gnani agent. The demonstrated call was started from the Agent Console."),
        ("Prisma ASR, Timbre 2.5, Evon LLM", "PARTIAL", "Prisma v2.5 product and Evon v2.0 Fast are evidenced. The agent uses Timbre G v1.0 because it was the only TTS model available in this Agent Console account; Timbre 2.5 was not offered."),
        ("Meaningful multi-turn conversation", "PASS", "Live recording and transcript captured before Agent Console credits were exhausted."),
        ("Post-call trigger received", "PARTIAL", "Simulated authenticated webhook received and stored; exact automatic provider webhook contract not captured."),
        ("Stage code and reason stored", "PASS", "MongoDB and dashboard evidence."),
        ("Outcome visible on dashboard", "PASS", "Scenario overview and completed Viji medical-hardship detail."),
        ("Invalid, failed trigger, duplicate handled", "PASS", "Explicit handlers and automated tests."),
        ("Runnable using README", "PASS", "Live Render URLs lead; localhost instructions are optional."),
    ]
    add_matrix(doc, ["Criterion", "Status", "Assessment"], [2820, 900, 5640], acceptance, status_index=1)

    doc.add_heading("6. Why I used mock mode", level=1)
    doc.add_paragraph(
        "The FDE document describes an outbound call-trigger API that receives customer data and invokes Gnani. My available Gnani Agent Console account did not expose a trigger endpoint, provider API key, or body template. The Initial Message callback I observed instead sent session metadata with blank phone fields and omitted customer and EMI variables."
    )
    add_kv_table(
        doc,
        [
            ("Observed callback fields", "call_id, flow_id, organization_id, environment, user_id, sender_id, blank phone_number/mobile"),
            ("Missing business fields", "customer_name, customer_id, loan account, EMI amount, due date, preferred language"),
            ("Credits limitation", "My Agent Console voice-call credits were exhausted; visible API-playground credits belong to a separate API product."),
            ("Demonstration fallback", "I used the FastAPI mock trigger to store the complete request/provider ID, then completed it with a simulated authenticated webhook."),
            ("Not claimed", "No claim that mock mode placed a phone call or called an unavailable Gnani trigger endpoint."),
        ],
    )
    add_callout(
        doc,
        "IMPORTANT DISTINCTION",
        "Pre-call variables configured inside the agent can guide its prompt, but the captured Initial Message request proves that those values were not serialized into this callback. A production integration must receive a documented Gnani trigger contract or resolve customer data from a trusted system using a stable identifier.",
        "partial",
    )

    doc.add_heading("7. Stage-code and idempotency logic", level=1)
    stage_rows = [
        ("PTP_TODAY", "Promise to pay today", "Promise to pay"),
        ("PTP_FUTURE", "Promise on future date", "Promise to pay"),
        ("ALREADY_PAID", "Payment already completed", "Paid"),
        ("CALLBACK_SCHEDULED", "Requested callback", "Follow-up"),
        ("RTP_FINANCIAL", "Refused due to financial hardship", "Other completed"),
        ("RTP_MEDICAL", "Unable to pay due to medical hardship", "Other completed"),
        ("DISPUTE_CHARGES", "EMI/charge dispute", "Other completed"),
        ("THIRD_PARTY", "Third party answered", "Other completed"),
        ("DSCN", "Disconnected without outcome", "Unreachable"),
        ("trigger_failed", "Provider trigger failed or timed out", "Failure"),
    ]
    add_matrix(doc, ["Stage code", "Meaning", "Dashboard group"], [2200, 4560, 2600], stage_rows)
    doc.add_paragraph(
        "Post-call delivery requires X-Webhook-API-Key. X-Webhook-Id is used as the delivery key; if omitted, the backend derives a SHA-256 fingerprint. The repository performs an atomic conditional update only when the delivery ID is absent. Repeating the same event returns duplicate and does not apply the outcome twice."
    )

    doc.add_heading("8. Submission requirement inventory", level=1)
    inventory = [
        ("Source code", "PASS", "GitHub repository"),
        ("Gnani configuration/export", "PASS", "Verified Agent Console configuration with screenshots covering agent details, languages, model, voice, analytics fields, and conversation logs"),
        ("Bot prompt and flow", "PASS", "Agent screenshots and repository documentation"),
        ("FastAPI application", "PASS", "backend/app"),
        ("Dummy dashboard", "PASS", "frontend plus live Render URL"),
        ("Database schema", "PASS", "README data model and MongoDB screenshot"),
        ("Postman/cURL", "PASS", "postman/CollectFlow.postman_collection.json and README"),
        (".env.example", "PASS", "Repository root"),
        ("Dockerfile", "PASS", "Backend and frontend Dockerfiles"),
        ("docker-compose.yml", "PASS", "Repository root"),
        ("README setup", "PASS", "Live deployment first; local setup optional"),
        ("Architecture diagram", "PASS", "This document and docs/architecture.md"),
        ("Sample recordings", "PASS", "submission/recordings MP3"),
        ("Sample webhook payloads", "PASS", "Postman collection, README, scenario runner"),
        ("Dashboard screenshots", "PASS", "submission/evidence/dashboard"),
        ("Stage-code logic", "PASS", "This document and code"),
        ("Mandatory test results", "PASS", "Dashboard overview plus automated runner"),
    ]
    add_matrix(doc, ["Requirement", "Status", "Location / evidence"], [3100, 900, 5360], inventory, status_index=1)

    doc.add_heading("9. Bonus requirements", level=1)
    doc.add_paragraph("I list only the bonus items that I passed or partially demonstrated.")
    bonus = [
        ("Dockerised deployment", "PASS", "Two Dockerfiles and Compose definition"),
        ("Call analytics charts", "PARTIAL", "Summary KPI cards exist; charting is not implemented"),
        ("API authentication", "PARTIAL", "Webhook key authentication implemented; dashboard/operator auth not implemented"),
        ("PII masking", "PASS", "Phone numbers are masked on dashboard rows and detail headers"),
        ("Detailed audit logs", "PARTIAL", "Structured logs and webhook IDs; full audit ledger not implemented"),
        ("Cloud deployment", "PASS", "Render plus MongoDB Atlas"),
        ("CI/CD automated tests", "PARTIAL", "Automated tests exist and Render deploys from Git; no captured CI workflow evidence"),
    ]
    add_matrix(doc, ["Bonus item", "Status", "Assessment"], [3100, 900, 5360], bonus, status_index=1)

    doc.add_heading("10. Security and privacy", level=1)
    for text in (
        "Secrets remain in environment variables and are excluded from Git.",
        "Post-call webhooks require a constant-time compared API key.",
        "Delivery IDs make webhook retries idempotent and auditable.",
        "The frontend never receives MongoDB credentials.",
        "Dashboard phone numbers are masked; selected PDF screenshots avoid full customer phone numbers.",
        "Any webhook credential displayed during development should be rotated before submission.",
        "Demo records use synthetic customers and outcomes.",
    ):
        add_bullet(doc, text)

    doc.add_heading("11. Production-readiness plan", level=1)
    add_callout(
        doc,
        "FREE-TIER DEPLOYMENT OPERATION",
        "I currently use Render's free tier, so the API and dashboard can become inactive after an idle period. Before demonstrating the project, I deploy the latest commit to bring the services back to a fresh live state and then refresh the dashboard. The first backend request can take about 50 seconds while the free-tier service starts.",
        "info",
    )
    readiness = [
        ("Provider contract", "Obtain Gnani trigger URL, auth scheme, IP ranges, payload schema, retry behavior, and signed webhook contract."),
        ("Data enrichment", "Resolve customer/EMI data from the system of record using a stable ID; do not trust session-only callbacks."),
        ("Authentication", "Add operator identity, role-based access, session expiry, and least-privilege roles."),
        ("Webhook security", "Prefer signed payloads with replay window, rotate keys, and alert on auth failures."),
        ("Reliability", "Queue provider calls/webhooks, add exponential retry, dead-letter handling, and scheduled workers."),
        ("Observability", "Add correlation IDs, structured audit events, metrics, traces, alerts, and provider dashboards."),
        ("Privacy", "Encrypt sensitive fields, define retention/deletion, restrict Atlas access, and review logs for PII."),
        ("Quality", "Run backend, frontend, contract, and end-to-end tests in CI before deploy; use staging gates."),
        ("Scale", "Add pagination, indexes, rate limits, pool tuning, and backpressure for bulk campaigns."),
    ]
    add_matrix(doc, ["Area", "Recommended action"], [1900, 7460], readiness)

    doc.add_heading("12. Evidence appendix", level=1)
    doc.add_paragraph(
        "I retained the following figures under submission/evidence to make this PDF self-contained. The dashboard filtering evidence uses the corrected, working Promise-to-pay result."
    )

    figures = [
        (ROOT / "submission/evidence/test-results/01-mandatory-scenarios-dashboard-overview.png", "Figure 2. Mandatory scenario records visible on the deployed dashboard", "Nine distinct customer outcomes demonstrate the completed business scenarios."),
        (ROOT / "submission/evidence/dashboard/07-promise-to-pay-filter-passed.png", "Figure 3. Corrected Promise-to-pay dashboard filter", "The working filter returns three matching completed calls."),
        (ROOT / "submission/evidence/dashboard/03-viji-medical-completed.png", "Figure 4. Completed medical-hardship call detail", "Shows outcome summary, RTP_MEDICAL, duration, transcript, and verified webhook payload."),
        (ROOT / "submission/evidence/mongodb/01-completed-call-document.png", "Figure 5. MongoDB Atlas completed-call document", "The record includes stage code, webhook ID, transcript, outcome, reason, PTP date, and provider ID."),
        (ROOT / "submission/evidence/workflow/01-create-mock-call-form.png", "Figure 6. FastAPI mock-call initiation form", "The form is explicitly labeled as mock mode and states that no phone call is placed."),
        (ROOT / "submission/evidence/agent-console/01-agent-details-language-switch.png", "Figure 7. Gnani agent details and language-switch configuration", "English/Spanish, region, time zone, description, and switch threshold are visible."),
        (ROOT / "submission/evidence/agent-console/02-evon-llm-settings.png", "Figure 8. Gnani Evon LLM configuration", "The captured model is Gnani Evon v2.0 Fast."),
        (ROOT / "submission/evidence/agent-console/03-timbre-voice-settings.png", "Figure 9. Gnani TTS configuration", "The agent uses Timbre G v1.0 with Jenny because it was the only TTS model available in this Agent Console account; Timbre 2.5 was not offered."),
        (ROOT / "submission/evidence/agent-console/04-analytics-extraction-fields.png", "Figure 10. Structured post-call extraction fields", "Disposition reason, summary, PTP date, callback time, and language are configured."),
        (ROOT / "submission/evidence/agent-console/05-live-call-recording-transcript.png", "Figure 11. Real Gnani multi-turn voice-test recording and transcript", "Conversation 9b8a1885-efd3-4d9a-9fdc-367de019e116 lasted 1m08s."),
        (ROOT / "submission/evidence/limitations/01-gnani-api-credits-prisma.png", "Figure 12. Gnani Prisma API playground and separate API credits", "This proves Prisma v2.5 API access, not Agent Console call credits or agent transcriber selection."),
        (ROOT / "submission/evidence/render/01-fastapi-web-service-live.png", "Figure 13. FastAPI deployed as a live Render Web Service", "The backend is publicly reachable over HTTPS."),
        (ROOT / "submission/evidence/render/02-dashboard-static-site-live.png", "Figure 14. Dashboard deployed as a live Render Static Site", "The production Vite build completed and the site is live."),
    ]
    for path, caption, note in figures:
        add_figure(doc, path, caption, note)

    doc.add_page_break()
    doc.add_heading("13. Project summary and transparency", level=1)
    add_callout(
        doc,
        "CORE APPLICATION: BUILT AND DEMONSTRATED",
        "I implemented and evidenced FastAPI validation, pending-call persistence, post-call idempotency, structured outcome storage, MongoDB persistence, dashboard visibility, search/filter behavior, failure paths, Docker packaging, and cloud deployment.",
        "pass",
    )
    add_callout(
        doc,
        "PROVIDER INTEGRATION: DOCUMENTED LIMITATIONS",
        "I demonstrated a real Gnani call, conversation, and recording. I could not fully demonstrate provider-triggered outbound calling, exact automatic webhook delivery, or the agent Prisma transcriber selection. I used Timbre G v1.0 because it was the only TTS model available in my Agent Console account; Timbre 2.5 was not offered.",
        "partial",
    )
    doc.add_paragraph(
        "I created this project as a technically complete and reproducible application-layer demonstration. The deployed demo uses Render's free tier, so I deploy the latest commit and refresh the dashboard before a live demonstration when the services have been inactive. To complete the remaining provider-specific integration, Gnani must issue the outbound-trigger credentials and restore Agent Console call credits. I can then run one final live call using the same call_id through initiation, conversation, webhook, MongoDB update, and dashboard display."
    )
    add_callout(
        doc,
        "TRANSPARENCY NOTE",
        "The limitation is caused by unavailable trigger credentials and exhausted Agent Console credits, not hidden by the submission.",
        "partial",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
